#!/usr/bin/env python3
"""Generate professional Chinese resume in Word format."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    """Add a section heading with bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=12, bold=True, color=(31, 73, 125))
    else:
        set_font(run, size=11, bold=True, color=(31, 73, 125))
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_bullet(doc, text, indent=0.4, bold_prefix=None):
    """Add a bullet point; optionally bold the prefix up to first colon."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(indent)
    if bold_prefix and "：" in text:
        prefix, rest = text.split("：", 1)
        r1 = p.add_run(prefix + "：")
        set_font(r1, bold=True)
        r2 = p.add_run(rest)
        set_font(r2)
    else:
        run = p.add_run(text)
        set_font(run)
    return p

def add_normal(doc, text, bold=False, size=10.5, space_before=2, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def add_job_header(doc, company, title, period):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(company + " — " + title)
    set_font(r1, size=11, bold=True)
    r2 = p.add_run("    " + period)
    set_font(r2, size=10, color=(89, 89, 89))

def add_sub_section(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_font(run, size=10.5, bold=True, color=(68, 114, 196))

# ── Header ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run("乔德立")
set_font(r, size=20, bold=True, color=(31, 73, 125))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(2)
p2.paragraph_format.space_after  = Pt(4)
r2 = p2.add_run("15522013699  |  qiaodeli111@126.com  |  深圳  |  github.com/qiaodeli111")
set_font(r2, size=10, color=(89, 89, 89))

# ── 个人优势 ──────────────────────────────────────────────────────────────────
add_heading(doc, "个人优势")
summary = (
    "14年DevOps/SRE经验，专注基础设施自动化与平台工程。主导构建了覆盖物理机（BMAAS）、"
    "私有云（CloudStack）及多云（AWS/阿里云/华为云/天翼云）的全栈IaC体系。具备Go/Python"
    "二次开发能力，自研多云Terraform Provider及JetDev开发者工具箱。率先完成龙蜥Anolis OS "
    "ARM、UOS海光C86、OpenEuler等国产化平台全链路适配，支撑x86/ARM64/C86三架构标准化交付。"
    "持有 AWS SAA、CKA、RHCE、Azure Admin、PMP 认证。"
)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run(summary)
set_font(run, size=10.5)

# ── 核心技术栈 ────────────────────────────────────────────────────────────────
add_heading(doc, "核心技术栈")

table = doc.add_table(rows=10, cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

skill_data = [
    ("云平台",      "CloudStack · AWS · 阿里云 · 华为云 · 天翼云 · Proxmox VE"),
    ("容器编排",    "Kubernetes · Docker · Helm · Harbor · JFrog Artifactory"),
    ("IaC 自动化",  "Terraform（自研Provider） · Ansible · Packer · Tinkerbell · cloud-init"),
    ("存储",        "ZFS · DRBD · Linstor · LVM · Corosync/Pacemaker"),
    ("网络",        "L4/L7 LB · BGP · VLAN · OVS/DPDK · PXE/iPXE · Redfish"),
    ("监控可观测",  "Grafana · Prometheus · Loki · Fluentd · Zabbix"),
    ("安全合规",    "Vault · Nessus · Lynis · Consul ACL · SSL/TLS · Trivy"),
    ("CI/CD",       "GitLab CI · Jenkins · Goss · Ansible Tower"),
    ("操作系统",    "Ubuntu · 龙蜥 Anolis OS · UOS · OpenEuler · Windows（x86/ARM64/C86）"),
    ("开发语言",    "Go · Python · Shell/Bash · HCL"),
]

for i, (cat, skills) in enumerate(skill_data):
    row = table.rows[i]
    # Category cell
    c0 = row.cells[0]
    c0.width = Cm(3.0)
    cp = c0.paragraphs[0]
    cp.paragraph_format.space_before = Pt(1)
    cp.paragraph_format.space_after  = Pt(1)
    cr = cp.add_run(cat)
    set_font(cr, size=9.5, bold=True, color=(31, 73, 125))
    # Skills cell
    c1 = row.cells[1]
    sp = c1.paragraphs[0]
    sp.paragraph_format.space_before = Pt(1)
    sp.paragraph_format.space_after  = Pt(1)
    sr = sp.add_run(skills)
    set_font(sr, size=9.5)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── 工作经历 ──────────────────────────────────────────────────────────────────
add_heading(doc, "工作经历")

# ── 深圳捷誊 ──────────────────────────────────────────────────────────────────
add_job_header(doc,
    "深圳捷誊技术有限公司",
    "DevOps Team Lead & 交付效能专家",
    "2024.10 – 至今 | 深圳")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run(
    "同时管理 BMAAS、Terraform、镜像工厂、CloudStack、K8s、IT支撑、JetDev 共 7个并行项目，"
    "主导技术选型与架构设计，带领团队完成跨 5+ 客户站点的交付。"
)
set_font(run, size=10, italic=True, color=(89, 89, 89))

add_sub_section(doc, "▌ 平台化建设")
add_bullet(doc,
    "主导研发 JetDev 一站式DevOps工具箱：集成 Ansible/Terraform/Packer/GitLab Runner，"
    "Ansible Role 脚手架生成时间缩短 80%，建立从代码审查到线上部署的标准化可追溯工作流。",
    bold_prefix="主导研发 JetDev 一站式DevOps工具箱")
add_bullet(doc,
    "自研多云统一 Terraform Provider（Go）：兼容 AWS/阿里云/华为云/天翼云，"
    "基于 AWS 兼容 API 实现 CloudStack 资源编排，消除多云配置碎片化，实现 IaaS 资源[一次编写、到处运行]。",
    bold_prefix="自研多云统一 Terraform Provider（Go）")
add_bullet(doc,
    "构建 CloudStack 私有云平台（ZFS+Linstor+DRBD 9.3.0 + VLAN高级网络）：完成 13项核心功能交付；"
    "ZFS 性能调优后 IOPS 从 230k 提升至近 800k（PCIe 5.0 32核）；修复 DRBD 心跳参数实现 7×24 稳定运行；"
    "上线 Grafana 资源看板实时展示 CPU/内存/存储/公网IP分配率。",
    bold_prefix="构建 CloudStack 私有云平台")

add_sub_section(doc, "▌ 自动化体系")
add_bullet(doc,
    "重构自动化镜像工厂（Packer）：支持 Ubuntu/UOS/龙蜥Anolis OS/OpenEuler/Windows 共 5类OS × "
    "x86_64/ARM64/C86 3架构全矩阵；集成 Goss 自动化冒烟测试 + Lynis 安全审计；并发构建效率提升 3.2倍，"
    "流水线零失败连续运行超 30天，完成 12项镜像交付任务。",
    bold_prefix="重构自动化镜像工厂（Packer）")
add_bullet(doc,
    "深度优化 BMAAS 裸金属自动装机平台：新增磁盘自动选择、多网卡 Bonding、HookOS 可插拔驱动框架、"
    "Rescue 模式及离线装机流程；支持 x86/ARM64/C86 全架构；成功交付扬州电信新站点 OS 全流程部署；"
    "上线 redfish-batch-iso 工具实现批量 ISO 并发挂载。",
    bold_prefix="深度优化 BMAAS 裸金属自动装机平台")
add_bullet(doc,
    "设计并落地多套 GitLab CI/CD 流水线：镜像工厂自动触发构建/上传/通知全链路；"
    "Ansible 项目[提交即测试]（动态VM部署+幂等性验证）；IaC 资源变更 MR 审批流。",
    bold_prefix="设计并落地多套 GitLab CI/CD 流水线")

add_sub_section(doc, "▌ 国产化适配与性能工程")
add_bullet(doc,
    "主导国产化平台全链路适配：完成龙蜥 Anolis OS 8 ARM JetIce 镜像打包（含 ZFS/DRBD/"
    "Corosync/Pacemaker HA 联调）；适配 UOS（海光C86）及 OpenEuler ZFS 模块编译验证。",
    bold_prefix="主导国产化平台全链路适配")
add_bullet(doc,
    "主导存储性能工程：设计并执行 fio 基准测试套件（随机/顺序读写+混合负载 6类场景，PCIe 3.0/5.0对比），"
    "输出可归因性能报告；完成 ZFS thin-provisioning 评估与 Linstor 副本迁移策略验证。",
    bold_prefix="主导存储性能工程")
add_bullet(doc,
    "完成 L4/L7 负载均衡性能测试：执行完整测试用例，汇总数据并输出跨团队评审性能报告；"
    "完成 L4LB Docker 容器化迁移及 BGP 动态路由管理。",
    bold_prefix="完成 L4/L7 负载均衡性能测试")

add_sub_section(doc, "▌ 客户交付与技术支撑")
add_bullet(doc,
    "主导多站点客户交付与现场问题排查：空天院网络丢包根因定位（网卡驱动/内核参数/DPDK/OVS）；"
    "厦门电信 6项高危 CVE 修复复测通过（含 CVE-2023-27350）；吉山数字生活 L4LB 容器化+"
    "Consul ACL 安全基线全站达标；四川攀枝花巡检工具重部署；无锡 Zabbix 全量接入凌霄平台。",
    bold_prefix="主导多站点客户交付与现场问题排查")
add_bullet(doc,
    "推进 K8s 集群平台化：验证 Cluster API + CloudStack 方案，确定 K8s on CloudStack + 外部LB + "
    "共享存储核心架构；规划 Velero+MinIO 备份与灾难演练；JumpServer Helm Chart 迁移 K8s + SSL 证书升级。",
    bold_prefix="推进 K8s 集群平台化")
add_bullet(doc,
    "主导公司内部网络整改：IP 盘点与归属映射、VLAN 划分、DNS/DHCP 策略部署，完成 19项 IT 基础设施交付任务。",
    bold_prefix="主导公司内部网络整改")

# ── Kyndryl / IBM ─────────────────────────────────────────────────────────────
add_job_header(doc,
    "Kyndryl 深圳（IBM拆分） / IBM 深圳",
    "SRE 工程师（工具链方向）",
    "2017.03 – 2024.08 | 深圳")

add_bullet(doc, "连续 6年 年度考核[超出预期]（Exceed Expectation），管理 4-10人 跨文化团队（中/印）。")
add_bullet(doc,
    "构建全面自动化运维体系：运维效率提升 50%+，质量提升 90%+，应用部署时间缩短 80%，"
    "变更请求处理时间减少 50%+。")
add_bullet(doc,
    "管理 Chef/GitLab/Jenkins/Ansible Tower/UCD 全链路 DevOps 工具链，"
    "为业务 SRE 团队开发 Ansible/Chef 公共基础库。")
add_bullet(doc,
    "作为 AWS 架构咨询团队成员，主导客户全球化架构改造：ECS 单体→EKS 集群化部署，"
    "引入 AWS Global Accelerator 分流加速，分离存储/数据库层，解决跨区访问延迟与成本问题。")

# ── 天津恩恩 ──────────────────────────────────────────────────────────────────
add_job_header(doc,
    "天津恩恩科技有限公司",
    "变更服务交付经理",
    "2016.05 – 2017.03")

add_bullet(doc,
    "重构变更管理流程：Excel模板 + 自动化脚本，团队工作量减少 80%，月度报表自动化生成。")

# ── DXC ───────────────────────────────────────────────────────────────────────
add_job_header(doc,
    "DXC Technology",
    "WebHosting 团队队长 / 运维工程师",
    "2012.06 – 2016.05")

add_bullet(doc,
    "带领 4人国际团队（中/印）为亚洲多地 AIA 友邦保险提供中间件运维，"
    "自动化程度从 10% 提升至 50%。")
add_bullet(doc,
    "主导数据中心迁移、网站分割、新系统建立等复杂项目，成功完成 IBM WAS6 → WAS8 迁移。")

# ── 认证证书 ──────────────────────────────────────────────────────────────────
add_heading(doc, "认证证书")
certs = (
    "AWS SAA（2022.11） · Azure Admin（2022.04） · 阿里云ACP（2021.07） · CKA（2021.08） · "
    "RHCE Ansible 2.9/RHEL 8（2022.06） · PMP（2017.11） · ITIL® Foundation（2017.02） · "
    "IBM WAS Network Deployment V8.0（2014.08）"
)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run(certs)
set_font(run, size=10)

# ── 教育背景 ──────────────────────────────────────────────────────────────────
add_heading(doc, "教育背景")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(2)
r1 = p.add_run("天津工业大学（双一流）")
set_font(r1, bold=True)
r2 = p.add_run("  |  软件工程 本科  |  2008 – 2012")
set_font(r2)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "/home/deliqiao/jetdev/github/ResumeSkills/output/resume_zh.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
